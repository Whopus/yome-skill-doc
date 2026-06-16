#!/usr/bin/env python3
"""Python/python-docx backend for @yome/doc on headless hosts.

The CLI runner talks to this script over JSON stdin/stdout:
  python doc_backend.py --probe
  python doc_backend.py --dispatch

This backend intentionally handles file-backed .docx operations directly and
keeps only lightweight active-document state under ~/.yome/state.
"""

from __future__ import annotations

import argparse
import json
import os
import re
import shutil
import sys
import time
from pathlib import Path
from typing import Any, Callable, Dict, List, Optional, Tuple

try:
    import docx
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_COLOR_INDEX
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.shared import Pt, RGBColor
except Exception as exc:  # pragma: no cover - exercised through --probe
    docx = None  # type: ignore[assignment]
    _IMPORT_ERROR = exc
else:
    _IMPORT_ERROR = None


BACKEND_ID = "python-docx"
STATE_ROOT = Path(os.environ.get("YOME_STATE_HOME") or (Path.home() / ".yome" / "state"))
STATE_FILE = STATE_ROOT / "@yome" / "doc" / "linux-session.json"
SUPPORTED_ACTIONS = [
    "open",
    "new",
    "save",
    "close",
    "files",
    "get",
    "read",
    "stats",
    "set",
    "append",
    "insert",
    "delete",
    "fmt",
    "find",
    "replace",
    "replace.all",
    "highlight",
    "table.add",
    "table.delete",
    "table.row.add",
    "table.row.delete",
    "table.col.add",
    "table.col.delete",
    "table.cell.set",
    "table.merge",
    "header.set",
    "footer.set",
    "header.clear",
    "footer.clear",
    "break.page",
    "break.line",
    "sections",
    "link.add",
    "image.add",
    "style.apply",
    "styles",
    "list.bullet",
    "list.number",
    "export.txt",
    "version",
    "screen",
    "alerts",
]

COLOR_NAMES = {
    "black": (0, 0, 0),
    "white": (255, 255, 255),
    "red": (255, 0, 0),
    "green": (0, 128, 0),
    "blue": (0, 0, 255),
    "yellow": (255, 255, 0),
    "gray": (128, 128, 128),
    "grey": (128, 128, 128),
    "orange": (255, 165, 0),
    "purple": (128, 0, 128),
    "pink": (255, 192, 203),
}

HIGHLIGHT_NAMES = {
    "yellow": WD_COLOR_INDEX.YELLOW if _IMPORT_ERROR is None else None,
    "green": WD_COLOR_INDEX.BRIGHT_GREEN if _IMPORT_ERROR is None else None,
    "pink": WD_COLOR_INDEX.PINK if _IMPORT_ERROR is None else None,
    "blue": WD_COLOR_INDEX.TURQUOISE if _IMPORT_ERROR is None else None,
    "red": WD_COLOR_INDEX.RED if _IMPORT_ERROR is None else None,
    "gray": WD_COLOR_INDEX.GRAY_25 if _IMPORT_ERROR is None else None,
    "grey": WD_COLOR_INDEX.GRAY_25 if _IMPORT_ERROR is None else None,
}

ALIGN_NAMES = {
    "left": WD_ALIGN_PARAGRAPH.LEFT if _IMPORT_ERROR is None else None,
    "center": WD_ALIGN_PARAGRAPH.CENTER if _IMPORT_ERROR is None else None,
    "centre": WD_ALIGN_PARAGRAPH.CENTER if _IMPORT_ERROR is None else None,
    "right": WD_ALIGN_PARAGRAPH.RIGHT if _IMPORT_ERROR is None else None,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY if _IMPORT_ERROR is None else None,
}


def main() -> int:
    parser = argparse.ArgumentParser(add_help=False)
    parser.add_argument("--probe", action="store_true")
    parser.add_argument("--dispatch", action="store_true")
    ns = parser.parse_args()

    if ns.probe:
        return emit_probe()
    if ns.dispatch:
        if _IMPORT_ERROR is not None:
            return emit_result(False, stderr=dependency_error(), exit_code=127)
        try:
            payload = json.load(sys.stdin)
            result = dispatch(payload)
            return emit_result(**result)
        except Exception as exc:
            return emit_result(False, stderr="doc python backend: {0}".format(exc), exit_code=1)

    return emit_result(False, stderr="usage: doc_backend.py --probe | --dispatch", exit_code=2)


def emit_probe() -> int:
    if _IMPORT_ERROR is not None:
        return emit_json({"ok": False, "backend": BACKEND_ID, "stderr": dependency_error()})
    return emit_json(
        {
            "ok": True,
            "backend": BACKEND_ID,
            "engine": "python-docx",
            "python_docx": getattr(docx, "__version__", "unknown"),
            "supports": SUPPORTED_ACTIONS,
        }
    )


def emit_result(ok: bool, stdout: str = "", stderr: str = "", exit_code: Optional[int] = None, **extra: Any) -> int:
    code = exit_code if exit_code is not None else (0 if ok else 1)
    return emit_json({"ok": ok, "stdout": stdout, "stderr": stderr, "exitCode": code, **extra}, code)


def emit_json(obj: Dict[str, Any], exit_code: int = 0) -> int:
    sys.stdout.write(json.dumps(obj, ensure_ascii=False))
    sys.stdout.write("\n")
    return exit_code


def dependency_error() -> str:
    return (
        "python-docx backend requires python-docx. Install it with "
        "`python3 -m pip install python-docx`, or set YOME_PYTHON to a Python "
        "environment that has python-docx. Import error: {0}".format(_IMPORT_ERROR)
    )


def dispatch(req: Dict[str, Any]) -> Dict[str, Any]:
    action = str(req.get("action") or "")
    positionals = [str(v) for v in req.get("positionals") or []]
    flags = {str(k): v for k, v in (req.get("flags") or {}).items()}
    cwd = str(req.get("workingDirectory") or os.environ.get("YOME_WORKING_DIRECTORY") or os.getcwd())
    state = load_state()

    if action == "open":
        return doc_open(positionals, flags, cwd, state)
    if action == "new":
        return doc_new(positionals, flags, cwd, state)
    if action == "save":
        return doc_save(positionals, flags, cwd, state)
    if action == "close":
        return doc_close(flags, state)
    if action == "files":
        return doc_files(state)
    if action == "version":
        return ok_json({"backend": BACKEND_ID, "engine": "python-docx", "python_docx": getattr(docx, "__version__", "unknown")})
    if action in ("screen", "alerts"):
        return ok("{0} ignored by headless {1}".format(action, BACKEND_ID))

    handlers: Dict[str, Callable[[Any, str, List[str], Dict[str, Any], Dict[str, Any], str], Dict[str, Any]]] = {
        "get": doc_get,
        "read": doc_read,
        "stats": doc_stats,
        "set": doc_set,
        "append": doc_append,
        "insert": doc_insert,
        "delete": doc_delete,
        "fmt": doc_fmt,
        "find": doc_find,
        "replace": doc_replace,
        "replace.all": doc_replace_all,
        "highlight": doc_highlight,
        "table.add": doc_table_add,
        "table.delete": doc_table_delete,
        "table.row.add": doc_table_row_add,
        "table.row.delete": doc_table_row_delete,
        "table.col.add": doc_table_col_add,
        "table.col.delete": doc_table_col_delete,
        "table.cell.set": doc_table_cell_set,
        "table.merge": doc_table_merge,
        "header.set": doc_header_set,
        "footer.set": doc_footer_set,
        "header.clear": doc_header_clear,
        "footer.clear": doc_footer_clear,
        "break.page": doc_break_page,
        "break.line": doc_break_line,
        "sections": doc_sections,
        "link.add": doc_link_add,
        "image.add": doc_image_add,
        "style.apply": doc_style_apply,
        "styles": doc_styles,
        "list.bullet": doc_list_bullet,
        "list.number": doc_list_number,
        "export.txt": doc_export_txt,
    }

    handler = handlers.get(action)
    if handler is None:
        return fail(
            "doc {0}: not supported by {1}. Supported: {2}".format(
                action, BACKEND_ID, ", ".join(SUPPORTED_ACTIONS)
            ),
            127,
        )
    return with_document(state, flags, cwd, positionals, handler)


def doc_open(positionals: List[str], flags: Dict[str, Any], cwd: str, state: Dict[str, Any]) -> Dict[str, Any]:
    raw = first(positionals) or str(flags.get("path") or "")
    path = resolve_user_path(raw, cwd)
    if not path:
        return fail("doc open: missing <path>", 2)
    if Path(path).suffix.lower() == ".doc":
        return fail("doc open: .doc is not supported by python-docx backend; use .docx", 1)
    if not Path(path).exists():
        return fail("doc open: file not found: {0}".format(path), 1)
    document = Document(path)
    remember_open(state, path)
    save_state(state)
    return ok_json({"ok": True, "opened": path, "paragraphs": len(document.paragraphs), "tables": len(document.tables)})


def doc_new(positionals: List[str], flags: Dict[str, Any], cwd: str, state: Dict[str, Any]) -> Dict[str, Any]:
    raw = first(positionals) or str(flags.get("path") or "")
    path = resolve_user_path(raw, cwd) if raw else str(default_new_path())
    if Path(path).suffix.lower() != ".docx":
        path = path + ".docx"
    force = truthy(flags.get("force"))
    if Path(path).exists() and not force:
        return fail("doc new: {0} already exists; pass --force=true to overwrite".format(path), 1)
    ensure_parent(path)
    document = Document()
    document.save(path)
    remember_open(state, path)
    save_state(state)
    return ok_json({"ok": True, "created": path})


def doc_save(positionals: List[str], flags: Dict[str, Any], cwd: str, state: Dict[str, Any]) -> Dict[str, Any]:
    src = active_path(state, flags, cwd)
    if not src:
        return fail("doc save: no active document", 1)
    raw_dest = str(flags.get("path") or first(positionals) or "")
    if not raw_dest:
        return ok_json({"ok": True, "path": src, "note": "document changes are saved after each write"})
    dest = resolve_user_path(raw_dest, cwd)
    if not dest:
        return fail("doc save: invalid --path", 2)
    if Path(dest).suffix.lower() != ".docx":
        return fail("doc save: python-docx backend can save-as only to .docx", 1)
    if Path(dest).exists() and dest != src and not truthy(flags.get("force")):
        return fail("doc save: {0} already exists; pass --force=true to overwrite".format(dest), 1)
    ensure_parent(dest)
    shutil.copyfile(src, dest)
    remember_open(state, dest)
    save_state(state)
    return ok_json({"ok": True, "path": dest})


def doc_close(flags: Dict[str, Any], state: Dict[str, Any]) -> Dict[str, Any]:
    path = state.get("activePath")
    if not path:
        return fail("doc close: no active document", 1)
    paths = [p for p in state.get("openPaths", []) if p != path]
    state["openPaths"] = paths
    state["activePath"] = paths[0] if paths else None
    save_state(state)
    return ok_json({"ok": True, "closed": path, "saved": flags.get("save") != "false"})


def doc_files(state: Dict[str, Any]) -> Dict[str, Any]:
    paths = [str(p) for p in state.get("openPaths", []) if Path(str(p)).exists()]
    state["openPaths"] = paths
    if state.get("activePath") not in paths:
        state["activePath"] = paths[0] if paths else None
    save_state(state)
    rows = ["index\tactive\tpath"]
    active = state.get("activePath")
    for idx, path in enumerate(paths, 1):
        rows.append("{0}\t{1}\t{2}".format(idx, "yes" if path == active else "no", path))
    return ok("\n".join(rows))


def doc_get(document: Any, path: str, positionals: List[str], flags: Dict[str, Any], state: Dict[str, Any], cwd: str) -> Dict[str, Any]:
    rows = ["index\ttext"]
    for idx, paragraph in enumerate(document.paragraphs, 1):
        rows.append("{0}\t{1}".format(idx, tsv_cell(preview(paragraph.text, 80))))
    return ok("\n".join(rows))


def doc_read(document: Any, path: str, positionals: List[str], flags: Dict[str, Any], state: Dict[str, Any], cwd: str) -> Dict[str, Any]:
    idx = parse_index(first(positionals) or flags.get("index"), "doc read")
    paragraph = paragraph_at(document, idx)
    run = first_run(paragraph)
    bold = run.bold if run is not None else None
    italic = run.italic if run is not None else None
    size = run.font.size.pt if run is not None and run.font.size is not None else ""
    name = run.font.name if run is not None and run.font.name is not None else ""
    color = ""
    if run is not None and run.font.color is not None and run.font.color.rgb is not None:
        color = "#{0}".format(run.font.color.rgb)
    return ok(
        "{0}\t{1}\t{2}\t{3}\t{4}\t{5}".format(
            tsv_cell(paragraph.text),
            bool_word(bold),
            bool_word(italic),
            size,
            tsv_cell(name),
            color,
        )
    )


def doc_stats(document: Any, path: str, positionals: List[str], flags: Dict[str, Any], state: Dict[str, Any], cwd: str) -> Dict[str, Any]:
    text = document_text(document)
    words = len(re.findall(r"\S+", text))
    paragraphs = len(document.paragraphs)
    chars = len(text)
    pages = 0
    return ok("{0}\t{1}\t{2}\t{3}".format(words, paragraphs, chars, pages))


def doc_set(document: Any, path: str, positionals: List[str], flags: Dict[str, Any], state: Dict[str, Any], cwd: str) -> Dict[str, Any]:
    idx = parse_index(first(positionals) or flags.get("index"), "doc set")
    text = required_text(positionals[1:] if len(positionals) > 1 else [], flags, "doc set")
    paragraph_at(document, idx).text = text
    document.save(path)
    return ok("set {0}".format(idx))


def doc_append(document: Any, path: str, positionals: List[str], flags: Dict[str, Any], state: Dict[str, Any], cwd: str) -> Dict[str, Any]:
    text = required_text(positionals, flags, "doc append")
    document.add_paragraph(text)
    document.save(path)
    return ok("appended")


def doc_insert(document: Any, path: str, positionals: List[str], flags: Dict[str, Any], state: Dict[str, Any], cwd: str) -> Dict[str, Any]:
    idx = parse_index(first(positionals) or flags.get("index"), "doc insert")
    text = required_text(positionals[1:] if len(positionals) > 1 else [], flags, "doc insert")
    paragraph_at(document, idx).insert_paragraph_before(text)
    document.save(path)
    return ok("inserted {0}".format(idx))


def doc_delete(document: Any, path: str, positionals: List[str], flags: Dict[str, Any], state: Dict[str, Any], cwd: str) -> Dict[str, Any]:
    idx = parse_index(first(positionals) or flags.get("index"), "doc delete")
    paragraph = paragraph_at(document, idx)
    element = paragraph._element
    element.getparent().remove(element)
    document.save(path)
    return ok("deleted {0}".format(idx))


def doc_fmt(document: Any, path: str, positionals: List[str], flags: Dict[str, Any], state: Dict[str, Any], cwd: str) -> Dict[str, Any]:
    idx = parse_index(first(positionals) or flags.get("index"), "doc fmt")
    paragraph = paragraph_at(document, idx)
    ensure_runs(paragraph)
    if "align" in flags:
        align = str(flags.get("align") or "").lower()
        if align not in ALIGN_NAMES:
            return fail("doc fmt: unsupported --align={0}".format(align), 2)
        paragraph.alignment = ALIGN_NAMES[align]
    for run in paragraph.runs:
        if "bold" in flags:
            run.bold = truthy(flags.get("bold"))
        if "italic" in flags:
            run.italic = truthy(flags.get("italic"))
        if "size" in flags:
            run.font.size = Pt(float(str(flags.get("size"))))
        if "font" in flags:
            run.font.name = str(flags.get("font") or "")
        if "color" in flags:
            rgb = parse_color(str(flags.get("color") or ""))
            run.font.color.rgb = RGBColor(rgb[0], rgb[1], rgb[2])
    document.save(path)
    return ok("formatted {0}".format(idx))


def doc_find(document: Any, path: str, positionals: List[str], flags: Dict[str, Any], state: Dict[str, Any], cwd: str) -> Dict[str, Any]:
    what = first(positionals) or str(flags.get("what") or "")
    if not what:
        return fail("doc find: missing <keyword>", 2)
    pattern = compile_search(what, flags)
    for idx, paragraph in enumerate(document.paragraphs, 1):
        match = pattern.search(paragraph.text)
        if match:
            return ok("found at paragraph {0} offset {1}..{2}".format(idx, match.start(), match.end()))
    return ok("not found")


def doc_replace(document: Any, path: str, positionals: List[str], flags: Dict[str, Any], state: Dict[str, Any], cwd: str) -> Dict[str, Any]:
    count = replace_text(document, positionals, flags, replace_all=False)
    if count:
        document.save(path)
        return ok("replaced 1")
    return ok("no match")


def doc_replace_all(document: Any, path: str, positionals: List[str], flags: Dict[str, Any], state: Dict[str, Any], cwd: str) -> Dict[str, Any]:
    count = replace_text(document, positionals, flags, replace_all=True)
    if count:
        document.save(path)
    return ok("replaced {0}".format(count))


def doc_highlight(document: Any, path: str, positionals: List[str], flags: Dict[str, Any], state: Dict[str, Any], cwd: str) -> Dict[str, Any]:
    idx = parse_index(first(positionals) or flags.get("index"), "doc highlight")
    color = str(flags.get("color") or "yellow").lower()
    paragraph = paragraph_at(document, idx)
    ensure_runs(paragraph)
    value = None if color == "none" else HIGHLIGHT_NAMES.get(color)
    if color != "none" and value is None:
        return fail("doc highlight: unsupported --color={0}".format(color), 2)
    for run in paragraph.runs:
        run.font.highlight_color = value
    document.save(path)
    return ok("highlighted {0}".format(idx))


def doc_table_add(document: Any, path: str, positionals: List[str], flags: Dict[str, Any], state: Dict[str, Any], cwd: str) -> Dict[str, Any]:
    rows = int(str(flags.get("rows") or "3"))
    cols = int(str(flags.get("cols") or "3"))
    table = document.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    if flags.get("index") is not None:
        idx = parse_index(flags.get("index"), "doc table.add")
        anchor = paragraph_at(document, idx)
        anchor._element.addnext(table._element)
    document.save(path)
    return ok_json({"ok": True, "table": len(document.tables), "rows": rows, "cols": cols})


def doc_table_delete(document: Any, path: str, positionals: List[str], flags: Dict[str, Any], state: Dict[str, Any], cwd: str) -> Dict[str, Any]:
    table = table_at(document, parse_index(first(positionals) or flags.get("index"), "doc table.delete"))
    element = table._element
    element.getparent().remove(element)
    document.save(path)
    return ok("deleted table")


def doc_table_row_add(document: Any, path: str, positionals: List[str], flags: Dict[str, Any], state: Dict[str, Any], cwd: str) -> Dict[str, Any]:
    table = table_at(document, parse_index(first(positionals) or flags.get("table"), "doc table.row.add"))
    count = int(str(flags.get("count") or "1"))
    for _ in range(count):
        table.add_row()
    document.save(path)
    return ok("added {0} row(s)".format(count))


def doc_table_row_delete(document: Any, path: str, positionals: List[str], flags: Dict[str, Any], state: Dict[str, Any], cwd: str) -> Dict[str, Any]:
    table = table_at(document, parse_index(first(positionals) or flags.get("table"), "doc table.row.delete"))
    row_idx = parse_index(flags.get("row"), "doc table.row.delete --row")
    if row_idx < 1 or row_idx > len(table.rows):
        raise ValueError("doc table.row.delete: row index out of range")
    table._tbl.remove(table.rows[row_idx - 1]._tr)
    document.save(path)
    return ok("deleted row {0}".format(row_idx))


def doc_table_col_add(document: Any, path: str, positionals: List[str], flags: Dict[str, Any], state: Dict[str, Any], cwd: str) -> Dict[str, Any]:
    table = table_at(document, parse_index(first(positionals) or flags.get("table"), "doc table.col.add"))
    count = int(str(flags.get("count") or "1"))
    for _ in range(count):
        table.add_column(Pt(72))
    document.save(path)
    return ok("added {0} col(s)".format(count))


def doc_table_col_delete(document: Any, path: str, positionals: List[str], flags: Dict[str, Any], state: Dict[str, Any], cwd: str) -> Dict[str, Any]:
    table = table_at(document, parse_index(first(positionals) or flags.get("table"), "doc table.col.delete"))
    col_idx = parse_index(flags.get("col"), "doc table.col.delete --col")
    for row in table.rows:
        cells = row.cells
        if col_idx < 1 or col_idx > len(cells):
            raise ValueError("doc table.col.delete: col index out of range")
        row._tr.remove(cells[col_idx - 1]._tc)
    document.save(path)
    return ok("deleted col {0}".format(col_idx))


def doc_table_cell_set(document: Any, path: str, positionals: List[str], flags: Dict[str, Any], state: Dict[str, Any], cwd: str) -> Dict[str, Any]:
    table = table_at(document, parse_index(first(positionals) or flags.get("table"), "doc table.cell.set"))
    row = parse_index(flags.get("row"), "doc table.cell.set --row")
    col = parse_index(flags.get("col"), "doc table.cell.set --col")
    text = required_text([], flags, "doc table.cell.set")
    table.cell(row - 1, col - 1).text = text
    document.save(path)
    return ok("set table cell {0},{1}".format(row, col))


def doc_table_merge(document: Any, path: str, positionals: List[str], flags: Dict[str, Any], state: Dict[str, Any], cwd: str) -> Dict[str, Any]:
    table = table_at(document, parse_index(first(positionals) or flags.get("table"), "doc table.merge"))
    from_row = parse_index(flags.get("fromRow") or flags.get("from_row"), "doc table.merge --fromRow")
    from_col = parse_index(flags.get("fromCol") or flags.get("from_col"), "doc table.merge --fromCol")
    to_row = parse_index(flags.get("toRow") or flags.get("to_row"), "doc table.merge --toRow")
    to_col = parse_index(flags.get("toCol") or flags.get("to_col"), "doc table.merge --toCol")
    table.cell(from_row - 1, from_col - 1).merge(table.cell(to_row - 1, to_col - 1))
    document.save(path)
    return ok("merged table cells")


def doc_header_set(document: Any, path: str, positionals: List[str], flags: Dict[str, Any], state: Dict[str, Any], cwd: str) -> Dict[str, Any]:
    text = required_text(positionals, flags, "doc header.set")
    for section in document.sections:
        set_story_text(section.header, text)
    document.save(path)
    return ok("header set")


def doc_footer_set(document: Any, path: str, positionals: List[str], flags: Dict[str, Any], state: Dict[str, Any], cwd: str) -> Dict[str, Any]:
    text = required_text(positionals, flags, "doc footer.set")
    for section in document.sections:
        set_story_text(section.footer, text)
    document.save(path)
    return ok("footer set")


def doc_header_clear(document: Any, path: str, positionals: List[str], flags: Dict[str, Any], state: Dict[str, Any], cwd: str) -> Dict[str, Any]:
    for section in document.sections:
        set_story_text(section.header, "")
    document.save(path)
    return ok("header cleared")


def doc_footer_clear(document: Any, path: str, positionals: List[str], flags: Dict[str, Any], state: Dict[str, Any], cwd: str) -> Dict[str, Any]:
    for section in document.sections:
        set_story_text(section.footer, "")
    document.save(path)
    return ok("footer cleared")


def doc_break_page(document: Any, path: str, positionals: List[str], flags: Dict[str, Any], state: Dict[str, Any], cwd: str) -> Dict[str, Any]:
    paragraph = break_target(document, flags)
    paragraph.add_run().add_break(WD_BREAK.PAGE)
    document.save(path)
    return ok("page break inserted")


def doc_break_line(document: Any, path: str, positionals: List[str], flags: Dict[str, Any], state: Dict[str, Any], cwd: str) -> Dict[str, Any]:
    paragraph = break_target(document, flags)
    paragraph.add_run().add_break(WD_BREAK.LINE)
    document.save(path)
    return ok("line break inserted")


def doc_sections(document: Any, path: str, positionals: List[str], flags: Dict[str, Any], state: Dict[str, Any], cwd: str) -> Dict[str, Any]:
    rows = ["index\tstart_type"]
    for idx, section in enumerate(document.sections, 1):
        rows.append("{0}\t{1}".format(idx, getattr(section.start_type, "name", str(section.start_type))))
    return ok("\n".join(rows))


def doc_link_add(document: Any, path: str, positionals: List[str], flags: Dict[str, Any], state: Dict[str, Any], cwd: str) -> Dict[str, Any]:
    idx = parse_index(first(positionals) or flags.get("index"), "doc link.add")
    url = str(flags.get("url") or "")
    if not url:
        return fail("doc link.add: missing --url", 2)
    text = str(flags.get("text") or url)
    paragraph = paragraph_at(document, idx)
    add_hyperlink(paragraph, text, url)
    document.save(path)
    return ok("link added")


def doc_image_add(document: Any, path: str, positionals: List[str], flags: Dict[str, Any], state: Dict[str, Any], cwd: str) -> Dict[str, Any]:
    raw = first(positionals) or str(flags.get("path") or "")
    image_path = resolve_user_path(raw, cwd)
    if not image_path:
        return fail("doc image.add: missing <image_path>", 2)
    if not Path(image_path).exists():
        return fail("doc image.add: file not found: {0}".format(image_path), 1)
    paragraph = document.add_paragraph()
    paragraph.add_run().add_picture(image_path)
    document.save(path)
    return ok("image added")


def doc_style_apply(document: Any, path: str, positionals: List[str], flags: Dict[str, Any], state: Dict[str, Any], cwd: str) -> Dict[str, Any]:
    idx = parse_index(first(positionals) or flags.get("index"), "doc style.apply")
    style = str(flags.get("style") or "")
    if not style:
        return fail("doc style.apply: missing --style", 2)
    paragraph_at(document, idx).style = style
    document.save(path)
    return ok("style applied")


def doc_styles(document: Any, path: str, positionals: List[str], flags: Dict[str, Any], state: Dict[str, Any], cwd: str) -> Dict[str, Any]:
    rows = ["name\ttype"]
    for style in document.styles:
        rows.append("{0}\t{1}".format(tsv_cell(style.name), getattr(style.type, "name", str(style.type))))
    return ok("\n".join(rows))


def doc_list_bullet(document: Any, path: str, positionals: List[str], flags: Dict[str, Any], state: Dict[str, Any], cwd: str) -> Dict[str, Any]:
    idx = parse_index(first(positionals) or flags.get("index"), "doc list.bullet")
    paragraph_at(document, idx).style = "List Bullet"
    document.save(path)
    return ok("bullet list applied")


def doc_list_number(document: Any, path: str, positionals: List[str], flags: Dict[str, Any], state: Dict[str, Any], cwd: str) -> Dict[str, Any]:
    idx = parse_index(first(positionals) or flags.get("index"), "doc list.number")
    paragraph_at(document, idx).style = "List Number"
    document.save(path)
    return ok("number list applied")


def doc_export_txt(document: Any, path: str, positionals: List[str], flags: Dict[str, Any], state: Dict[str, Any], cwd: str) -> Dict[str, Any]:
    raw = first(positionals) or str(flags.get("path") or "")
    dest = resolve_user_path(raw, cwd)
    if not dest:
        return fail("doc export.txt: missing <path>", 2)
    if Path(dest).exists() and not truthy(flags.get("force")):
        return fail("doc export.txt: {0} already exists; pass --force=true to overwrite".format(dest), 1)
    ensure_parent(dest)
    Path(dest).write_text(document_text(document), encoding="utf-8")
    return ok("exported")


def with_document(
    state: Dict[str, Any],
    flags: Dict[str, Any],
    cwd: str,
    positionals: List[str],
    handler: Callable[[Any, str, List[str], Dict[str, Any], Dict[str, Any], str], Dict[str, Any]],
) -> Dict[str, Any]:
    path = active_path(state, flags, cwd)
    if not path:
        return fail("doc: no active document; run `doc open <path>` or `doc new <path>` first", 1)
    if Path(path).suffix.lower() == ".doc":
        return fail("doc: .doc is not supported by python-docx backend; use .docx", 1)
    if not Path(path).exists():
        return fail("doc: active document not found: {0}".format(path), 1)
    document = Document(path)
    return handler(document, path, positionals, flags, state, cwd)


def active_path(state: Dict[str, Any], flags: Dict[str, Any], cwd: str) -> str:
    raw = flags.get("file") or flags.get("doc") or flags.get("document")
    if raw:
        return resolve_user_path(str(raw), cwd)
    active = state.get("activePath")
    return str(active or "")


def load_state() -> Dict[str, Any]:
    try:
        if STATE_FILE.exists():
            data = json.loads(STATE_FILE.read_text(encoding="utf-8"))
            if isinstance(data, dict):
                data.setdefault("openPaths", [])
                return data
    except Exception:
        pass
    return {"openPaths": [], "activePath": None}


def save_state(state: Dict[str, Any]) -> None:
    STATE_FILE.parent.mkdir(parents=True, exist_ok=True)
    STATE_FILE.write_text(json.dumps(state, ensure_ascii=False, indent=2), encoding="utf-8")


def remember_open(state: Dict[str, Any], path: str) -> None:
    paths = [str(p) for p in state.get("openPaths", []) if str(p) != path and Path(str(p)).exists()]
    paths.insert(0, path)
    state["openPaths"] = paths
    state["activePath"] = path


def resolve_user_path(raw: str, cwd: str) -> str:
    if not raw:
        return ""
    raw = os.path.expanduser(raw)
    path = Path(raw)
    if not path.is_absolute():
        path = Path(cwd) / path
    return str(path.resolve())


def default_new_path() -> Path:
    return Path.cwd() / "yome-doc-{0}.docx".format(int(time.time() * 1000))


def ensure_parent(path: str) -> None:
    Path(path).parent.mkdir(parents=True, exist_ok=True)


def ok(stdout: str = "") -> Dict[str, Any]:
    return {"ok": True, "stdout": stdout}


def ok_json(obj: Dict[str, Any]) -> Dict[str, Any]:
    return ok(json.dumps(obj, ensure_ascii=False))


def fail(stderr: str, exit_code: int = 1) -> Dict[str, Any]:
    return {"ok": False, "stdout": "", "stderr": stderr, "exit_code": exit_code}


def first(values: List[Any]) -> str:
    return str(values[0]) if values else ""


def truthy(value: Any) -> bool:
    if isinstance(value, bool):
        return value
    if value is None:
        return False
    return str(value).strip().lower() not in ("", "0", "false", "no", "off")


def parse_index(value: Any, label: str) -> int:
    if value is None or str(value) == "":
        raise ValueError("{0}: missing index".format(label))
    idx = int(str(value))
    if idx < 1:
        raise ValueError("{0}: index must be 1-based".format(label))
    return idx


def paragraph_at(document: Any, idx: int) -> Any:
    if idx < 1 or idx > len(document.paragraphs):
        raise ValueError("paragraph index out of range: {0}".format(idx))
    return document.paragraphs[idx - 1]


def table_at(document: Any, idx: int) -> Any:
    if idx < 1 or idx > len(document.tables):
        raise ValueError("table index out of range: {0}".format(idx))
    return document.tables[idx - 1]


def required_text(positionals: List[str], flags: Dict[str, Any], label: str) -> str:
    value = flags.get("text")
    if value is None and "with" in flags:
        value = flags.get("with")
    if value is None and positionals:
        value = first(positionals)
    if value is None:
        raise ValueError("{0}: missing --text".format(label))
    return str(value)


def tsv_cell(value: str) -> str:
    return value.replace("\t", " ").replace("\r", " ").replace("\n", " ")


def preview(value: str, limit: int) -> str:
    return value if len(value) <= limit else value[:limit] + "..."


def bool_word(value: Optional[bool]) -> str:
    if value is True:
        return "true"
    if value is False:
        return "false"
    return "mixed"


def first_run(paragraph: Any) -> Any:
    return paragraph.runs[0] if paragraph.runs else None


def ensure_runs(paragraph: Any) -> None:
    if paragraph.runs:
        return
    text = paragraph.text
    paragraph.text = ""
    paragraph.add_run(text)


def parse_color(value: str) -> Tuple[int, int, int]:
    v = value.strip().lower()
    if v in COLOR_NAMES:
        return COLOR_NAMES[v]
    if v.startswith("#") and len(v) == 7:
        return (int(v[1:3], 16), int(v[3:5], 16), int(v[5:7], 16))
    if "," in v:
        parts = [int(p.strip()) for p in v.split(",")]
        if len(parts) == 3 and all(0 <= p <= 255 for p in parts):
            return (parts[0], parts[1], parts[2])
    raise ValueError("invalid color: {0}".format(value))


def compile_search(what: str, flags: Dict[str, Any]) -> Any:
    escaped = re.escape(what)
    if truthy(flags.get("wholeWord")):
        escaped = r"\b{0}\b".format(escaped)
    regex_flags = 0 if truthy(flags.get("matchCase")) else re.IGNORECASE
    return re.compile(escaped, regex_flags)


def replace_text(document: Any, positionals: List[str], flags: Dict[str, Any], replace_all: bool) -> int:
    what = first(positionals) or str(flags.get("what") or "")
    replacement = str(flags.get("with") or "")
    if not what:
        raise ValueError("doc replace: missing <keyword>")
    pattern = compile_search(what, flags)
    total = 0
    max_count = 0 if replace_all else 1
    for paragraph in document.paragraphs:
        if max_count and total >= max_count:
            break
        count_for_para = 0 if replace_all else 1
        new_text, count = pattern.subn(replacement, paragraph.text, count=count_for_para)
        if count:
            paragraph.text = new_text
            total += count
    return total


def document_text(document: Any) -> str:
    parts: List[str] = []
    for paragraph in document.paragraphs:
        parts.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def set_story_text(story: Any, text: str) -> None:
    if story.paragraphs:
        story.paragraphs[0].text = text
        for paragraph in story.paragraphs[1:]:
            element = paragraph._element
            element.getparent().remove(element)
    else:
        story.add_paragraph(text)


def break_target(document: Any, flags: Dict[str, Any]) -> Any:
    if flags.get("index") is not None:
        return paragraph_at(document, parse_index(flags.get("index"), "doc break"))
    return document.add_paragraph()


def add_hyperlink(paragraph: Any, text: str, url: str) -> None:
    rel_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)

    run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")
    style = OxmlElement("w:rStyle")
    style.set(qn("w:val"), "Hyperlink")
    run_props.append(style)
    run.append(run_props)

    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


if __name__ == "__main__":
    raise SystemExit(main())
